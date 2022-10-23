import os, webbrowser, sys, requests, subprocess, pyttsx3
import smtplib
from email.mime.text import MIMEText
from pyfiglet import Figlet
import folium
import openpyxl
import mimetypes
from email import encoders
from translate import Translator
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
from email.mime.audio import MIMEAudio
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from docx import Document
from docx.shared import Inches, Mm
from openpyxl import load_workbook
import random
from win10toast import ToastNotifier
import time

'''
Инициализируем речь бота при старте программы
'''
engine = pyttsx3.init()
engine.setProperty('rate', 180)  # скорость речи


def speaker(text):
    engine.say(text)
    engine.runAndWait()


def browser():
    webbrowser.open('https://www.google.com/', new=2)


def youtube():
    webbrowser.open('https://www.youtube.com/', new=2)


def gmail():
    webbrowser.open('https://mail.ru/?from=logout', new=2)


def github():
    webbrowser.open('https://github.com/', new=2)


def python():
    webbrowser.open('https://docs.python.org/3/', new=2)


def pes():
    webbrowser.open('https://www.youtube.com/watch?v=ZQUuB-iVOpw', new=2)


def moment():
    webbrowser.open('https://youtu.be/GysF-VKy1hA?t=47', new=2)


def kartinka(mes):      # функция закидывания картинки на почту
    sender = "alekseiika74@gmail.com"
    password = "tohnylgqrvuojqer"
    server = smtplib.SMTP("smtp.gmail.com", 587)
    server.starttls()

    try:
        with open("email_template.html") as file:
            template = file.read()
    except IOError:
        return "Не найден документ"

    try:
        server.login(sender, password)
        msg = MIMEText(template, "html")
        msg = MIMEMultipart()
        msg["From"] = sender
        msg["To"] = "aleksei.grigorev74@mail.ru"
        msg["Subject"] = "Всё для меня любимого"

        msg.attach(MIMEText(mes))
        msg.attach(MIMEText(template, "html"))

        for file in os.listdir("attachmens"):
            filename = os.path.basename(file)
            ftype, encoding = mimetypes.guess_type(file)
            file_type, subtype = ftype.split("/")

            if file_type == "text":
                with open(f"attachmens/{file}") as f:
                    file = MIMEText(f.read())
            elif file_type == 'image':
                with open(f"attachmens/{file}", 'rb') as f:
                    file = MIMEImage(f.read(), subtype)
            elif file_type == 'audio':
                with open(f"attachmens/{file}", 'rb') as f:
                    file = MIMEAudio(f.read(), subtype)
            elif file_type == 'application':
                with open(f"attachmens/{file}", 'rb') as f:
                    file = MIMEApplication(f.read(), subtype)
            else:
                with open(f"attachmens/{file}", 'rb') as f:
                    file = MIMEBase(file_type, subtype)
                    file.set_payload(f.read())
                    encoders.encode_base64(file)

            file.add_header('content-disposition', 'attachment', filename=filename)
            msg.attach(file)
        server.sendmail(sender, "aleksei.grigorev74@mail.ru", msg.as_string())

        return speaker("Я успешно отправила сообщение")
    except Exception as _ex:
        return f"{_ex}\n Check login and password"


def prg():      # функция открытия программы exe
    try:
        speaker('Успешно открыто')
        # subprocess.Popen('')  # путь до exe
    except:
        speaker('Путь к файлу не найден')


def weather():      # функция погоды
    try:
        params = {'q': 'Saint Petersburg', 'units': 'metric', 'lang': 'ru', 'appid': 'be589f2249d9cd9b20dc2e19b11508d3'}
        response = requests.get(f'https://api.openweathermap.org/data/2.5/weather', params=params)
        if not response:        # чтобы не проверять все коды
            raise
        w = response.json()
        speaker(f"На улице {w['weather'][0]['description']} {round(w['main']['temp'])} градусов")
    except:
        speaker('Произошла ошибка при попытке запроса к ресурсу API, проверь код')


def offBot():       # функция выхода из бота
    sys.exit()


def passive():      # функция для диалогов
    pass


#   h53C7PvKzpfJ2MwNfRPp  пароль
# "alekseiika74@gmail.com"  почта
def send_email(message):        # функция потравки сообщения на почту
    print(message)
    sender = "alekseiika74@gmail.com"
    password = "tohnylgqrvuojqer"
    server = smtplib.SMTP("smtp.gmail.com", 587)
    server.starttls()

    try:
        server.login(sender, password)
        msg = MIMEText(message)
        msg["Subject"] = "Сообщение от Лёхи"
        server.sendmail(sender, "aleksei.grigorev74@mail.ru", msg.as_string())

        return speaker("Я успешно отправила сообщение")
    except Exception as _ex:
        return f"{_ex}\n Check login and password"


def get_info_ip(ip='5.18.244.153'):     # функция пробивания по ip
    speaker(ip)
    preview_text = Figlet(font='slant')
    print(preview_text.renderText('IP INFO'))

    # speaker('Введите адрес в консоль')
    # ip = input()
    # get_info_ip(ip)
    try:
        response = requests.get(url=f'http://ip-api.com/json/{ip}').json()
        #   print(response)
        slovar = {
            '[IP]': response.get('query'),
            '[Int prov]': response.get('isp'),
            '[Org]': response.get('org'),
            '[Country]': response.get('country'),
            '[Region Name]': response.get('regionName'),
            '[City]': response.get('city'),
            '[ZIP]': response.get('zip'),
            '[Lat]': response.get('lat'),
            '[Lon]': response.get('lon'),
        }

        for k, v in slovar.items():
            print(f'{k}:{v}')

        int_prov = response.get('isp')
        organization = response.get('org')
        country = response.get('country')
        region = response.get('regionName')
        city = response.get('city')
        ind = response.get('zip'),
        dolgota = response.get('lat')
        shirina = response.get('lon')

        translator = Translator(from_lang="en", to_lang="ru")
        to_int_prov = translator.translate(int_prov)
        to_organization = translator.translate(organization)
        to_country = translator.translate(country)
        to_city = translator.translate(city)

        if response.get('regionName') == 'St.-Petersburg':
            region = '78'

        speaker('Его интернет-провайдер' + ' ' + to_int_prov)
        speaker('Организация машрутизатора' + ' ' + to_organization)
        speaker('Страна' + ' ' + to_country)
        speaker('Регион' + ' ' + region)
        speaker('Город' + ' ' + to_city)
        speaker('Почтовый индекс')
        dl = len(ind)
        while dl > 0:
            for elem in str(ind):
                speaker(elem)
                dl -= 1
        speaker('Нашла его координаты')
        speaker('Долгота')
        dl_dolg = len(str(dolgota))
        while dl_dolg > 0:
            for elem in str(dolgota):
                if elem == '.':
                    speaker('точка')
                else:
                    speaker(elem)
                dl_dolg -= 1
        speaker('Широта')
        dl_shir = len(str(shirina))
        while dl_shir > 0:
            for elem in str(shirina):
                if elem == '.':
                    speaker('точка')
                else:
                    speaker(elem)
                dl_shir -= 1
        speaker('Начинаю анализ данных и поиск на карте')
        speaker('Вывожу карту в браузер')
        speaker('Я также вывела все данные в консоль и сохранила в ваш проект файл с картой')

        area = folium.Map(location=[response.get('lat'), response.get('lon')])
        area.save(f'{response.get("query")}_{response.get("city")}.html')
        webbrowser.open(
            'http://localhost:63342/Croxa/5.18.244.153_St%20Petersburg.html?_ijt=qaia0mv80lora09nq9i7407ori&_ij_reload=RELOAD_ON_SAVE')
    except requests.exceptions.ConnectionError:
        speaker('Возникли проблемы с подключением')


def create_doc():       # функция создания ворд документа
    speaker('Начала сборрку документа')
    documnt = Document()
    documnt.add_heading('Мотоциклетные шлема', 0)
    documnt.add_paragraph('Таблица шлемов')
    headers = ('№ п/п', 'Наименование товара', 'Стоимость', 'Наличие')
    rows = (
        (1, 'Спортивный шлем', 22799, 'есть'),
        (2, 'Эндуро шлем', 8999, 'нет'),
        (3, 'Кроссовый шлем', 5999, 'есть')
    )
    cols_number = len(headers)
    table = documnt.add_table(rows=1, cols=cols_number)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]
    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])
    documnt.add_picture('мото.jpg', width=Mm(50))
    documnt.save('testing.docx')
    speaker('Документ успешно собран, открываю')
    os.system('start testing.docx')


def zashita():      # функция создания заметки
    try:
        with open('zametki.txt', 'w') as file:
            file.write('Защтить курсовой проект\n')
        os.system('start zametki.txt')
    except:
        speaker('Ошибка при работе с файлом')


def dop():      # функция дополнения файла
    try:
        with open('zametki.txt', 'a') as file:
            file.write('Дожить до четвертой пары\n')
        os.system('start zametki.txt')
    except:
        speaker('Ошибка при работе с файлом')


def chitka_doc():       # функция читки документа
    document = Document('прочти.docx')
    paragraphs = document.paragraphs
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    speaker('\n'.join(text))
    speaker('Я закончила')


def excel():        # функция читки эксель файла
    wb = openpyxl.reader.excel.load_workbook(filename='data.xlsx')
    print(wb.sheetnames)
    wb.active = 0
    sheet = wb.active
    for i in range(1, 12):
        print(sheet['A' + str(i)].value, sheet['B' + str(i)].value, sheet['C' + str(i)].value)
    speaker('Я успешно загрузила данные в консоль')


def excel_random():     # функция заполнения рандомными данными
    list_a = ['шлем', 'перчатки', 'сковорода', 'пылесос', 'машина', 'ООП', 1, 40, -50, 2.654654, -2.2, 40.0943,
              'логистика', 654, 'эмпиризм', '00098', 'ключ', 'значение']
    fn = 'random.xlsx'
    wb = load_workbook(fn)
    ws = wb['random']

    for row in range(1, 4):
        for col in range(1, 4):
            value = random.choice(list_a)
            cell = ws.cell(row=row, column=col)
            cell.value = value
    wb.save(fn)
    wb.close()
    speaker('Задача выполнена, открываю файл')
    os.system('start random.xlsx')


def yved():     # функция уведомления
    zag = 'Сдать курсовик'
    con = 'Тебе нужно сдать его любой ценой'
    speaker('Я создала уведомление со следующими параметрами')
    speaker('Заголовок' + zag)
    speaker('c текстом' + con)
    speaker('Выведу через три секунды')
    time.sleep(3)
    speaker('Вывожу уведомление')
    toast = ToastNotifier()
    toast.show_toast(zag, con, duration=3, icon_path='timer.ico')


def date():     # функция выгрузка даты
    result = time.gmtime(time.time())
    mes = result.tm_mon
    day = result.tm_mday
    hour = result.tm_hour
    minuts = result.tm_min
    secynd = result.tm_sec
    dn = result.tm_wday
    mes_dict = {1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля', 5: 'мая', 6: 'июня', 7: 'июля',
                8: 'августа', 9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'}
    week_dict = {0: 'понедельник', 1: 'вторник', 2: 'среда', 3: 'четверг', 4: 'пятница', 5: 'суббота', 6: 'воскресенье'}
    hour_dict = {0: 'три ночи', 1: 'четыре утра', 2: 'пять утра', 3: 'шесть утра', 5: 'семь утра', 6: 'восемь утра',
                 7: 'девять утра', 8: 'десять утра', 9: 'одинадцать утра', 10: 'двенадцать дня', 11: 'час дня',
                 12: 'два часа дня', 13: 'три часа дня', 14: 'четыре часа дня', 15: 'пять вечера', 16: 'шесть вечера',
                 17: 'семь вечера', 18: 'восемь вечера', 19: 'девявть вечера', 20: 'десять вечера',
                 21: 'одинадцать вечера', 22: 'двенадцать ночи', 23: 'час ночи', 24: 'два ночи'}
    day_dict = {1: 'первое', 2: 'второе', 3: 'третье', 4: 'четвертое', 5: 'пятое', 6: 'шестое', 7: 'седьмое',
                8: 'восьмое', 9: 'девятое', 10: 'десятое', 11: 'одинадцатое', 12: 'двенадцатое', 13: 'тринадцатое',
                14: 'четырнадцатое', 15: 'пятнадцатое', 16: 'шестнадцатое', 17: 'семнадцатое', 18: 'восемнадцатое',
                19: 'девятнадцатое', 20: 'двадцатое', 21: 'двадцать первое', 22: 'двадцать второе',
                23: 'двадцать третье', 24: 'двадцать четвертое', 25: 'двадцать пятое', 26: 'двадцать шестое',
                27: 'двадцать седьмое', 28: 'двадцать восьмое', 29: 'двадцать девятое', 30: 'тридцатое',
                31: 'тридцать первое'}

    speaker('Сейчас две тысячи двадцать второй год')
    speaker(day_dict[day] + mes_dict[mes])
    speaker('Время по МСК' + hour_dict[hour] + str(minuts) + 'минуты' + str(secynd) + 'секунды')
    speaker('День недели' + week_dict[dn])
    speaker('Анализ времени закончен, приятного времяпровождения')


def fraza():        # функция рандомного ответа фразами
    fraz_list = ['вроде бы не плохо, но могла и лучше', 'более менее, думаю мне нужен отдых',
                 'просто великолепно, сегодня чудесный день для работы',
                 'замечательно, теперь во мне есть много функционала', 'я бы предпочла не отвечать на этот вопрос']

    value = random.choice(fraz_list)
    speaker(value)



